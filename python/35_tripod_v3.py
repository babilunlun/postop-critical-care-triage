#!/usr/bin/env python3
"""Build tripod_ai_checklist_v3.csv/.docx from v2 with INSPIRE temporal-validation updates."""
import csv
from docx import Document
from docx.shared import Pt

SRC = "/mnt/results/04_manuscript/tripod_ai_checklist_v2.csv"
CSV = "/mnt/results/04_manuscript/tripod_ai_checklist_v3.csv"

rows = list(csv.reader(open(SRC, encoding="utf-8")))
header, data = rows[0], rows[1:]

# item -> (Reported-in-manuscript location, Evidence/notes)
updates = {
    "2": ("Abstract",
          "149-word abstract covers setting, sample, model, outcome, internal performance, temporal "
          "(INSPIRE) and international (MOVER) external validation, recalibration, the GRU comparator, "
          "and the internal->temporal->external discrimination gradient."),
    "4": ("Introduction - Objectives",
          "Objectives: (i) develop the model; (ii) validate the frozen model in two independent cohorts "
          "\u2014 a same-institution earlier-decade cohort (INSPIRE; temporal validation) and a different-"
          "country cohort (MOVER; international external validation); (iii) head-to-head GRU vs XGBoost."),
    "5a": ("Methods - Data sources",
           "Three deidentified perioperative datasets described separately: VitalDB (development), "
           "INSPIRE v1.4.2 (same-institution temporal validation; PhysioNet), MOVER (international "
           "external validation)."),
    "5b": ("Methods - Data sources; Participants",
           "VitalDB Aug 2016\u2013Jun 2017; INSPIRE 2011\u20132020 (times recorded as minutes since each "
           "patient's first admission, no absolute dates); MOVER Nov 2017\u2013Aug 2023."),
    "6a": ("Methods - Data sources; Participants",
           "Development and temporal validation at Seoul National University Hospital (Korea); "
           "international validation at UC Irvine Medical Center (USA). INSPIRE shares the development "
           "institution but a different, earlier decade, isolating time/case-mix from site effects."),
    "6b": ("Methods - Participants",
           "Identical clinical gates applied to INSPIRE as to development (age >=18, general anesthesia, "
           "non-cardiac surgery, anesthesia duration >=30 min), yielding 98,633 operations before "
           "overlap exclusion."),
    "7": ("Methods - Predictors; Cohort overlap",
          "INSPIRE harmonized to the same 33 clinical + 38 time-series predictors: five-minute "
          "intraoperative series expanded to a one-minute grid (constant within epoch); drug/fluid "
          "exposures from the medication table (bolus sums + infusion integrals); implausible values "
          "set missing."),
    "8a": ("Methods - Outcome",
           "Composite of postoperative ICU admission or in-hospital death, defined per cohort; INSPIRE "
           "ICU flag = admission within 24 h of surgery end (slightly narrower window), disclosed as a "
           "limitation."),
    "9b": ("Methods - Predictors",
           "Predictors identically defined across cohorts; for INSPIRE, hypertension/diabetes "
           "ascertained by preoperative ICD-10 code (I10\u2013I16 / E10\u2013E14) OR an antihypertensive/"
           "antidiabetic prescription; rocuronium not recorded in INSPIRE and imputed by the pipeline."),
    "11": ("Methods - Missing data; Predictors",
           "The frozen imputation pipeline was applied unchanged to INSPIRE; unrecorded fields "
           "(e.g., rocuronium) and epochs without valid signal were handled as missing by the pipeline."),
    "12d": ("Methods - External validation; Results - Temporal/External/Subgroup/Sensitivity analyses",
            "Three-cohort design; performance heterogeneity quantified across internal, temporal "
            "(INSPIRE), and international (MOVER) validation, plus subgroup (department, year, patient "
            "class) and outcome-definition sensitivity subsets."),
    "12f": ("Methods - External validation and recalibration; Sensitivity analyses",
            "Intercept-only and Platt recalibration; 500x 50/50 split-sample validation applied in both "
            "INSPIRE and MOVER; recalibration also refit within each sensitivity subset."),
    "16": ("Methods - Data sources, Predictors, Cohort overlap; Discussion - Limitations",
           "Development-vs-INSPIRE evaluation differences documented: same institution/different decade; "
           "five-minute vs high-resolution signals (expanded to one-minute grid); ages in five-year "
           "bins; times relative to admission (no absolute dates); comorbidity from codes/prescriptions; "
           "rocuronium imputed; ICU window 24 h."),
    "17": ("Methods - Ethics",
           "INSPIRE released under SNUH IRB H-2210-078-1368 with waiver of consent, distributed under "
           "the Korea Credentialed Health Data License via PhysioNet; VitalDB and MOVER approvals also "
           "stated."),
    "18e": ("Methods - Data availability",
            "VitalDB at vitaldb.net; INSPIRE on PhysioNet (physionet.org/content/inspire/); MOVER under "
            "a data use agreement. Derived analysis datasets not redistributable under the MOVER "
            "agreement."),
    "20a": ("Results - Cohort characteristics; Fig. 1; Methods - Participants, Cohort overlap",
            "INSPIRE flow reported: 98,633 gated operations -> 2,437 development-linked exclusions -> "
            "96,196 operations (78,305 patients; 14,627 events, 15.2%); shown in the three-column Fig. 1 "
            "flow diagram."),
    "20b": ("Results - Cohort characteristics; Table 1",
            "Table 1 reports characteristics for all three cohorts (VitalDB, INSPIRE, MOVER) with "
            "per-cohort missingness and event-stratified p values."),
    "20c": ("Results - Cohort characteristics; Table 1; Discussion - Limitations",
            "Development (VitalDB) compared with both evaluation cohorts; comorbidity prevalence noted as "
            "not strictly comparable because of differing ascertainment across the three cohorts."),
    "23a": ("Tables 2; Figures 2\u20134; Supplementary Figures S2\u2013S3; Supplementary Tables S1\u2013S3, S5, S7",
            "AUROC/AUPRC/Brier/calibration with DeLong CIs for internal, INSPIRE, and MOVER; INSPIRE "
            "XGBoost AUROC 0.852 (0.848\u20130.855); ROC and calibration plots in Fig. 4."),
    "23b": ("Results - Temporal validation, Subgroup and robustness analyses; Discussion",
            "Discrimination gradient quantified: internal 0.932 -> temporal (INSPIRE) 0.852 -> "
            "international (MOVER) 0.794; calibration drift mild temporally, large internationally."),
    "24": ("Results - Temporal and External validation; Table 2; Figure 4",
           "Platt recalibration restored calibration in both INSPIRE (slope 1.000, Brier 0.092) and "
           "MOVER (slope 1.00), split-sample validated; INSPIRE pre-recalibration drift was mild "
           "(intercept -0.14, slope 0.71)."),
    "25": ("Discussion, para 1\u20134",
           "Nested validation gradient interpreted: most of the transportability penalty for absolute "
           "risk arises from crossing institutions and outcome definitions rather than time; the model "
           "stayed well calibrated across a decade within one institution but required recalibration "
           "abroad."),
    "26": ("Discussion - Limitations (sixth point)",
           "INSPIRE-specific limitations stated: same-institution (probes time/case-mix, not site); "
           "residual overlap bounded but not fully excluded (worst-case +0.0009 AUROC); five-year age "
           "bins; five-minute signal resolution; imputed rocuronium; differing comorbidity "
           "ascertainment; 24-h ICU window."),
}

n_applied = 0
for r in data:
    item = r[1]
    if item in updates:
        r[3], r[4] = updates[item]
        n_applied += 1

assert n_applied == len(updates), f"applied {n_applied} of {len(updates)} updates (item id mismatch?)"

with open(CSV, "w", encoding="utf-8", newline="") as f:
    w = csv.writer(f)
    w.writerow(header)
    w.writerows(data)
print("CSV saved:", CSV, "rows:", len(data), "updates applied:", n_applied)

doc = Document()
st = doc.styles["Normal"]
st.font.name = "Liberation Sans"
st.font.size = Pt(9)
doc.add_heading("TRIPOD+AI 2024 checklist - manuscript v3 (with INSPIRE temporal validation)", level=1)
t = doc.add_table(rows=len(data) + 1, cols=5)
t.style = "Table Grid"
for j, h in enumerate(header):
    t.cell(0, j).text = h
for i, r in enumerate(data, start=1):
    for j, val in enumerate(r):
        t.cell(i, j).text = val
for ri, row in enumerate(t.rows):
    for cell in row.cells:
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.size = Pt(8)
                run.font.name = "Liberation Sans"
                if ri == 0:
                    run.font.bold = True
doc.save("/workspace/tripod_ai_checklist_v3.docx")
print("DOCX saved")
