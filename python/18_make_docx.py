#!/usr/bin/env python3
"""Build manuscript_v1.docx from manuscript_v1.md + table CSVs (python-docx)."""
import csv
import re
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

MD = "/mnt/results/04_manuscript/manuscript_v1.md"
T1 = "/mnt/results/04_manuscript/table1_combined.csv"
T2 = "/mnt/results/04_manuscript/table2_performance.csv"
OUT = "/workspace/manuscript_v1.docx"

doc = Document()

# base style
style = doc.styles["Normal"]
style.font.name = "Liberation Sans"
style.font.size = Pt(11)
for lvl in ("Heading 1", "Heading 2"):
    st = doc.styles[lvl]
    st.font.name = "Liberation Sans"

text = open(MD, encoding="utf-8").read()
lines = text.split("\n")

def add_table_from_csv(path, doc, font_size=8):
    with open(path, encoding="utf-8") as f:
        rows = list(csv.reader(f))
    t = doc.add_table(rows=len(rows), cols=len(rows[0]))
    t.style = "Table Grid"
    for i, r in enumerate(rows):
        for j, val in enumerate(r):
            cell = t.cell(i, j)
            cell.text = val.strip()
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(font_size)
                    run.font.name = "Liberation Sans"
                    if i == 0:
                        run.font.bold = True
    return t

i = 0
in_refs = False
while i < len(lines):
    ln = lines[i].rstrip()
    i += 1
    if not ln.strip() or ln.strip() == "---":
        continue
    if ln.startswith("# "):  # title
        p = doc.add_heading(ln[2:].strip(), level=0)
        continue
    if ln.startswith("## "):
        h = ln[3:].strip()
        doc.add_heading(h, level=1)
        in_refs = (h == "References")
        # insert real tables at the Tables section
        if h == "Tables":
            doc.add_paragraph(
                "Table 1. Cohort characteristics. Values are mean (SD), median [IQR], or n (%). "
                "p values compare event vs no event within each cohort. Missing = percentage missing. "
                "V, VitalDB development cohort; M, MOVER external validation cohort. "
                "Emergency surgery was not recorded in MOVER (\u2014).")
            add_table_from_csv(T1, doc)
            doc.add_paragraph("")
            doc.add_paragraph(
                "Table 2. Model performance on internal and external validation. "
                "AUROC with DeLong 95% CI; AUPRC, area under the precision-recall curve; Cal, calibration. "
                "Recalibrated rows report split-sample-validated estimates (500 repetitions of 50/50 splits).")
            add_table_from_csv(T2, doc)
        continue
    if ln.startswith("### "):
        doc.add_heading(ln[4:].strip(), level=2)
        continue
    if ln.startswith("- "):
        doc.add_paragraph(ln[2:].strip(), style="List Bullet")
        continue
    # references: "1. ..." numbered
    m = re.match(r"^(\d+)\.\s+(.*)$", ln)
    if in_refs and m:
        doc.add_paragraph(f"{m.group(1)}. {m.group(2)}")
        continue
    # strip markdown emphasis/backticks for docx
    clean = ln.replace("**", "").replace("*", "").replace("`", "")
    clean = clean.replace("^", "")
    doc.add_paragraph(clean)

doc.save(OUT)
print("saved", OUT)
