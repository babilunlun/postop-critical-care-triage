#!/usr/bin/env python3
"""TRIPOD+AI 2024 checklist mapped to manuscript sections -> CSV + DOCX."""
import csv
from docx import Document
from docx.shared import Pt

ROWS = [
    # section, item, requirement (abbrev.), where reported, evidence/notes
    ("Title", "1", "Identify study as developing/evaluating a multivariable prediction model, target population, outcome",
     "Title page",
     "Title states development + external validation of an intraoperative ML model for postoperative critical care (14 words, no punctuation)."),
    ("Abstract", "2", "See TRIPOD+AI for Abstracts checklist",
     "Abstract",
     "132-word unstructured abstract covers setting, sample, model, outcome, internal/external performance, recalibration, DL comparator, conclusion."),
    ("Introduction - Background", "3a", "Healthcare context (diagnostic/prognostic) and rationale, with references to existing models",
     "Introduction, para 1-2",
     "Prognostic context (postoperative deterioration, ICU triage); SASA and prior ML/DL models cited [1-7]."),
    ("Introduction - Background", "3b", "Target population and intended purpose in care pathway, intended users",
     "Introduction, para 1; Discussion, para 5",
     "Adults under general anesthesia for non-cardiac surgery; end-of-surgery disposition aid for anesthesiologists/surgeons."),
    ("Introduction - Background", "3c", "Describe known health inequalities between sociodemographic groups",
     "Discussion, para 6 (partially)",
     "No formal inequality analysis; cohorts span Korean and US populations; flagged for revision if required."),
    ("Introduction - Objectives", "4", "Specify objectives; development or validation (or both)",
     "Introduction, para 3",
     "Three explicit objectives: development, external validation, head-to-head DL comparison."),
    ("Methods - Data", "5a", "Sources of data separately for development and evaluation; rationale; representativeness",
     "Methods - Data sources",
     "VitalDB (development; public high-fidelity registry) and MOVER (evaluation; public EHR+waveform repository) described separately."),
    ("Methods - Data", "5b", "Dates of data collection, start/end of accrual, end of follow-up",
     "Methods - Data sources",
     "VitalDB Aug 2016-Jun 2017; MOVER Nov 2017-Aug 2023."),
    ("Methods - Participants", "6a", "Study setting, number and location of centres",
     "Methods - Data sources / Participants",
     "Two tertiary academic centres: SNUH (Seoul, Korea) and UCI Medical Center (Orange County, USA)."),
    ("Methods - Participants", "6b", "Eligibility criteria",
     "Methods - Participants; Figure 1",
     "Adult (>=18 yr), general anesthesia, non-cardiac surgery, duration >=30 min; monitoring coverage >=50% (VitalDB)."),
    ("Methods - Participants", "6c", "Treatments received and handling during development/evaluation",
     "Methods - Predictors",
     "Intraoperative drug doses and fluid/blood volumes included as predictors; no treatment intervention studied."),
    ("Methods - Data preparation", "7", "Data pre-processing and quality checking",
     "Methods - Predictors / Missing data",
     "Unit harmonization, physiological plausibility caps, monitoring quality filter; identical pipeline applied to both cohorts."),
    ("Methods - Outcome", "8a", "Define outcome, time horizon, assessment, rationale, consistency across groups",
     "Methods - Outcome",
     "Composite of postoperative ICU admission or in-hospital death during index hospitalization; administrative flags; broader MOVER definition acknowledged."),
    ("Methods - Outcome", "8b", "If subjective interpretation, describe assessor qualifications",
     "Methods - Outcome",
     "Not applicable: outcome from administrative records without subjective interpretation."),
    ("Methods - Outcome", "8c", "Actions to blind outcome assessment",
     "Methods - Outcome",
     "Not applicable: routine administrative outcome; predictors measured before outcome."),
    ("Methods - Predictors", "9a", "Choice of initial predictors and any pre-selection",
     "Methods - Predictors",
     "All routinely available perioperative variables plus pre-specified engineered time-series summaries; no data-driven pre-selection."),
    ("Methods - Predictors", "9b", "Define all predictors, how/when measured, blinding",
     "Methods - Predictors; Supplementary Table S4",
     "33 clinical + 38 time-series predictors defined with measurement windows (preop labs <=30 d; intraop during anesthesia)."),
    ("Methods - Predictors", "9c", "If subjective predictor measurement, describe assessors",
     "Methods - Predictors",
     "Not applicable: device-recorded signals and structured EHR fields."),
    ("Methods - Sample size", "10", "How study size arrived at; justification",
     "Methods - Model development; Results",
     "All eligible cases used (development 5,987 with 1,184 events; evaluation 49,394 with 22,159 events); no formal calculation; events ample relative to tuned hyperparameters."),
    ("Methods - Missing data", "11", "How missing data were handled; reasons for omitting data",
     "Methods - Missing data",
     "Median imputation learned on training data; explicit EBL missingness indicator; MOVER laboratory missingness (56-69%) quantified in Table 1."),
    ("Methods - Analytical methods", "12a", "How data were used/partitioned for development and evaluation",
     "Methods - Model development",
     "Temporal 70/30 split (train 4,190 / test 1,797); 5-fold stratified CV within training for tuning; MOVER used only for evaluation."),
    ("Methods - Analytical methods", "12b", "How predictors were handled (functional form, transformation, standardisation)",
     "Methods - Missing data / Model development",
     "Median imputation, novel-level handling, one-hot encoding, zero-variance removal, normalization; engineered burden/trend features."),
    ("Methods - Analytical methods", "12c", "Type of model, rationale, model-building steps, hyperparameter tuning, internal validation",
     "Methods - Model development / Deep learning comparator",
     "SASA reference, LR (clinical/full), XGBoost (space-filling grid of 25, 5-fold CV); GRU architecture, training, early stopping, 3-seed ensemble."),
    ("Methods - Analytical methods", "12d", "Heterogeneity of parameter values/performance across clusters",
     "Methods - External validation; Results - Subgroup analyses",
     "Two-centre design; performance heterogeneity quantified by external validation and subgroup analyses (department, year, patient class)."),
    ("Methods - Analytical methods", "12e", "Measures and plots to evaluate performance and compare models",
     "Methods - Statistical analysis",
     "AUROC (DeLong CI), AUPRC, Brier, calibration intercept/slope, decision curves; DeLong test for model comparison."),
    ("Methods - Analytical methods", "12f", "Model updating (e.g., recalibration) arising from evaluation",
     "Methods - External validation and recalibration",
     "Intercept-only and Platt recalibration; 500x 50/50 split-sample validation for honest post-recalibration estimates."),
    ("Methods - Analytical methods", "12g", "How model predictions were calculated (formula, code, object, API)",
     "Methods - Model development; Code availability",
     "Frozen tidymodels workflow objects applied without refitting; code available on request/repository."),
    ("Methods - Class imbalance", "13", "If class imbalance methods used, state why/how and recalibration",
     "Methods - Model development",
     "No class-imbalance methods used (event rate 19.8%); not applicable."),
    ("Methods - Fairness", "14", "Approaches to address model fairness and rationale",
     "Results - Subgroup and robustness analyses",
     "No formal fairness framework; performance reported across department, patient-class, arterial-line, and calendar-year strata."),
    ("Methods - Model output", "15", "Output of the model (probabilities/classification); thresholds",
     "Methods - Statistical analysis",
     "Continuous probability output; no fixed classification threshold imposed; net benefit shown across thresholds (Fig. 2c)."),
    ("Methods - Training versus evaluation", "16", "Differences between development and evaluation data",
     "Methods - Data sources/Participants; Results - Cohort characteristics; Table 1",
     "Country, EHR, case mix, arterial-line practice, outcome definition, and missingness differences explicitly compared."),
    ("Methods - Ethical approval", "17", "Name IRB/ethics committee; consent or waiver",
     "Methods - Ethics",
     "SNUH IRB (VitalDB) and UCI IRB (MOVER) with data use agreement; secondary analysis of deidentified public data. [Authors to verify wording.]"),
    ("Open science - Funding", "18a", "Source of funding and role of funders",
     "Acknowledgements",
     "[To be completed by authors.]"),
    ("Open science - Conflicts of interest", "18b", "Declare conflicts of interest",
     "Competing interests",
     "Declared none. [Authors to confirm.]"),
    ("Open science - Protocol", "18c", "Where protocol can be accessed, or state none prepared",
     "Methods - Statistical analysis (to state)",
     "No formal protocol was prepared; analysis decisions pre-specified before external validation."),
    ("Open science - Registration", "18d", "Registration information, or state not registered",
     "Methods - Statistical analysis (to state)",
     "Not registered (retrospective analysis of public deidentified data)."),
    ("Open science - Data sharing", "18e", "Availability of study data",
     "Data availability",
     "VitalDB public; MOVER under DUA; derived data not redistributable."),
    ("Open science - Code sharing", "18f", "Availability of analytical code",
     "Code availability",
     "R code to be deposited in a public repository upon acceptance."),
    ("Patient & Public Involvement", "19", "Details of PPI or state no involvement",
     "(to state in manuscript)",
     "No patient or public involvement; retrospective analysis of deidentified public datasets."),
    ("Results - Participants", "20a", "Flow of participants with/without outcome; diagram",
     "Figure 1; Results - Cohort characteristics",
     "Full funnel with exclusions and event counts for both cohorts."),
    ("Results - Participants", "20b", "Characteristics overall and per data source; events; missing data",
     "Table 1; Results - Cohort characteristics",
     "Unified Table 1 with overall/event-stratified values and missingness for both cohorts."),
    ("Results - Participants", "20c", "Comparison of development vs evaluation data distributions",
     "Table 1; Results - Cohort characteristics",
     "Side-by-side VitalDB vs MOVER comparison including outcome rate and key predictors."),
    ("Results - Model development", "21", "Number of participants and outcome events in each analysis",
     "Figure 1; Results; Table 2",
     "Train 4,190 (823 events); test 1,797 (361); MOVER 49,394 (22,159); GRU subtrain/validation split reported."),
    ("Results - Model specification", "22", "Full prediction model details to allow predictions in new individuals",
     "Methods - Model development; Code availability",
     "Final hyperparameters reported; fitted R model object and prediction code available."),
    ("Results - Model performance", "23a", "Performance estimates with CIs, including key subgroups; plots",
     "Tables 2; Figures 2-4; Supplementary Tables S1-S3",
     "AUROC/AUPRC/Brier/calibration with DeLong CIs; subgroup AUROCs; ROC/calibration/decision-curve plots."),
    ("Results - Model performance", "23b", "Heterogeneity in performance across clusters",
     "Results - Subgroup and robustness analyses; Supplementary Table S2",
     "AUROC by department (0.769-0.825), year (0.780-0.807), patient class (0.722-0.787), arterial-line stratum."),
    ("Results - Model updating", "24", "Results of model updating and subsequent performance",
     "Results - External validation; Table 2; Figure 4",
     "Intercept-only and Platt recalibration with split-sample-validated Brier and calibration slope."),
    ("Discussion - Interpretation", "25", "Overall interpretation in context of objectives and previous studies",
     "Discussion, para 1-4",
     "Findings interpreted against prior single-center and DL studies; discrimination-vs-calibration message."),
    ("Discussion - Limitations", "26", "Limitations and effects on bias, uncertainty, generalizability",
     "Discussion, para 6",
     "Single-center retrospective development, broad MOVER outcome, informative lab missingness, BIS imputation, no prospective validation."),
    ("Discussion - Usability", "27a", "How poor quality/unavailable input data handled at implementation",
     "Methods - Missing data; Discussion, para 5-6",
     "Deployment uses the same imputation pipeline; external performance with 56-69% missing labs is a lower bound."),
    ("Discussion - Usability", "27b", "Required user interaction and expertise",
     "Discussion, para 5",
     "End-of-surgery decision aid for anesthesiologists/surgeons; not autonomous; no data manipulation by users."),
    ("Discussion - Usability", "27c", "Next steps for future research; applicability and generalizability",
     "Discussion, para 5-6",
     "Silent-mode prospective evaluation; local recalibration set; further validation cohorts."),
]

CSV = "/mnt/results/04_manuscript/tripod_ai_checklist_v1.csv"
with open(CSV, "w", newline="", encoding="utf-8") as f:
    w = csv.writer(f)
    w.writerow(["Section/Topic", "Item", "Checklist requirement (abbreviated)",
                "Reported in manuscript", "Evidence / notes"])
    w.writerows(ROWS)
print("CSV saved:", CSV, "rows:", len(ROWS))

doc = Document()
st = doc.styles["Normal"]; st.font.name = "Liberation Sans"; st.font.size = Pt(10)
doc.add_heading("TRIPOD+AI 2024 checklist", level=0)
doc.add_paragraph(
    "Manuscript: Development and external validation of an intraoperative machine "
    "learning model for postoperative critical care. Item wording abbreviated from "
    "Collins GS, et al. BMJ 2024;385:e078378.")
t = doc.add_table(rows=len(ROWS) + 1, cols=5)
t.style = "Table Grid"
hdr = ["Section/Topic", "Item", "Checklist requirement (abbreviated)",
       "Reported in manuscript", "Evidence / notes"]
for j, h in enumerate(hdr):
    c = t.cell(0, j); c.text = h
    for r in c.paragraphs[0].runs:
        r.font.bold = True; r.font.size = Pt(8)
for i, row in enumerate(ROWS, start=1):
    for j, val in enumerate(row):
        c = t.cell(i, j); c.text = val
        for p in c.paragraphs:
            for r in p.runs:
                r.font.size = Pt(8)
doc.save("/workspace/tripod_ai_checklist_v1.docx")
print("DOCX saved")
