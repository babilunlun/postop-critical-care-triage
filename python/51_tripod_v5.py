#!/usr/bin/env python3
"""Build tripod_ai_checklist_v5.csv/.docx from v4 with B1-B4 decision-impact updates."""
import csv
from docx import Document
from docx.shared import Pt

SRC = "/mnt/results/04_manuscript/tripod_ai_checklist_v4.csv"
CSV = "/mnt/results/04_manuscript/tripod_ai_checklist_v5.csv"
DOCX = "/workspace/tripod_ai_checklist_v5.docx"

rows = list(csv.reader(open(SRC, encoding="utf-8")))
header, data = rows[0], rows[1:]

updates = {
    "1": ("Title",
          "Title identifies development AND dual (temporal + international) external validation of a "
          "multivariable prediction model, the target population (non-cardiac surgery), the outcome "
          "(postoperative critical care need), and the ICU-course-verified triage evaluation."),
    "2": ("Abstract",
          "149-word abstract covers setting, sample, model, outcome, internal/temporal/international "
          "performance, the ICU-course decomposition (57.5%/72.8% observational admissions; "
          "admitted-patient discrimination AUROC 0.50), and the fixed 96%-sensitivity triage operating "
          "point (low-tier coverage 57.8%/28.0%, NPV 99.7%/98.3%, positive net benefit across 5-30% "
          "decision thresholds)."),
    "4": ("Introduction - Objectives",
          "Objectives: (i) develop the model; (ii) validate the frozen model in two independent cohorts "
          "(INSPIRE temporal; MOVER international); (iii) head-to-head GRU vs XGBoost; (iv) using "
          "ICU-course data, separate true critical care need from admission behavior and derive a "
          "risk-stratified triage pathway; (v) under a fixed safety constraint (>=96% sensitivity for "
          "true need), quantify decision impact - low-tier coverage, senior-review workload, and the "
          "observational bed-day review pool."),
    "12e": ("Methods - Statistical analysis",
            "AUROC (DeLong CI), AUPRC, Brier, calibration intercept/slope, decision curves (dcurves); "
            "DeLong test for model comparison. NEW: exact additive decomposition of the external "
            "calibration intercept into prevalence, case-mix, and residual components (development "
            "reference: VitalDB internal test set); subgroup transportability AUROCs with DeLong 95% "
            "CIs across pre-specified strata; decision-curve analysis of the need-recalibrated risk in "
            "both validation cohorts (thresholds 0.01-0.50)."),
    "15": ("Methods - Triage pathway evaluation; Results - A risk-stratified triage pathway",
           "Model output is a probability; pre-specified need-recalibrated risk bands (low <5%, "
           "intermediate 5-30%, high >30%) define the three-tier pathway, with percentile-band "
           "robustness checks. NEW: a safety-first operating point (highest need-recalibrated risk "
           "achieving >=96% sensitivity for true need; INSPIRE 2.0%, MOVER 4.6%) with per-1,000 "
           "decision-impact metrics (coverage, NPV, capture fractions, review workload, directly "
           "avoidable bed-days, observational review pool) and an equal-safety SASA comparator on the "
           "SASA-evaluable subset."),
    "14": ("Methods - Statistical analysis; Results - Subgroup and robustness analyses",
           "Formal fairness evaluation of the triage operating point across sex and age strata: "
           "stratum-specific sensitivity and low-tier NPV for true need at the cohort operating point "
           "(Clopper-Pearson 95% CIs), low-tier coverage, and calibration intercepts from offset "
           "logistic models; chi-square tests of sensitivity homogeneity. Sensitivity maintained at "
           "95.3-99.4% across all strata of both cohorts (target >=96%); low-tier NPV >97.7% "
           "throughout; coverage concentrated in younger/female patients mirroring lower baseline "
           "risk (Table S16, Fig. S6)."),
    "23a": ("Tables 2-5; Figures 2, 4-7; Supplementary Tables S1-S3, S5, S7-S16; Supplementary Figures S4-S6",
            "AUROC/AUPRC/Brier/calibration with DeLong CIs for internal, INSPIRE (0.907, 0.905-0.910), "
            "and de-duplicated MOVER (0.794, 0.790-0.798); refined-outcome AUROCs; triage-tier event "
            "rates (Table 4). NEW: decision-impact simulation at the >=96%-sensitivity operating point "
            "(Table 5, Fig. 7); validation-cohort DCA of need-recalibrated risk (Fig. 6, Table S12); "
            "calibration-intercept decomposition (Table S13, Fig. S4); subgroup transportability "
            "(Table S14, Fig. S5); equal-safety SASA comparator (Table S15); fairness of the "
            "operating point across sex and age strata (Table S16, Fig. S6)."),
    "23b": ("Results - Temporal/External validation, Subgroup and robustness analyses; Discussion",
            "Discrimination gradient: internal 0.932 -> temporal 0.907 -> international 0.794; "
            "calibration drift decomposed exactly: INSPIRE intercept -0.809 = prevalence -0.733 "
            "(90.5%) + case-mix +0.044 + residual -0.121; MOVER +2.352 = prevalence +1.178 (50.1%) + "
            "case-mix -0.191 + residual +1.365 (58.0%). NEW: subgroup need-discrimination ranges "
            "0.79-0.92 (INSPIRE) and 0.72-0.80 (MOVER), weakest in ASA 3-5 strata."),
    "24": ("Results - Temporal/External validation, Triage pathway; Tables 2, 4 and 5",
           "Platt recalibration restored calibration in INSPIRE (Brier 0.066, slope 0.999) and MOVER "
           "(Brier 0.183, slope 1.001); need-recalibrated tiers showed consistent safety (low-tier "
           "mortality <0.1%, NPV ~98%). NEW: at the >=96%-sensitivity operating point, low-tier NPV "
           "99.66% (INSPIRE) / 98.27% (MOVER); 94.2%/97.2% of missed escalations and 98.3%/99.6% of "
           "deaths captured above the threshold; equal-safety SASA coverage less than half the "
           "model's (20.2% vs 44.0%; 10.9% vs 27.7%)."),
    "25": ("Discussion, para 1-9",
           "Interpretation updated: discrimination transfers while absolute risk requires local "
           "recalibration; calibration-drift decomposition gives concrete deployment guidance "
           "(temporal: intercept-only update; geographic: full logistic recalibration or need "
           "redefinition); admission is not need; triage-aid deployment quantified at a fixed safety "
           "constraint - the model's advantage over SASA is efficiency at fixed safety, not safety "
           "itself; honest resource framing (directly avoidable bed-days modest at the development "
           "institution; larger observational review pool where admission is liberal)."),
    "26": ("Discussion - Limitations",
           "v4 limitations plus: decision-impact simulation is a retrospective what-if analysis under "
           "current practice (no prospective workflow change); directly avoidable bed-days estimable "
           "only in INSPIRE (MOVER lacks ICU length of stay); the 96% sensitivity constraint is "
           "cohort-specific, so operating-point thresholds (2.0%/4.6%) require local re-derivation; "
           "equal-safety SASA comparison is restricted to the SASA-evaluable subset."),
}

n_upd = 0
for r in data:
    item = r[1]
    if item in updates:
        r[3], r[4] = updates[item]
        n_upd += 1

with open(CSV, "w", newline="", encoding="utf-8") as f:
    w = csv.writer(f)
    w.writerow(header)
    w.writerows(data)
print(f"updated {n_upd} items -> {CSV}")

doc = Document()
st = doc.styles["Normal"]
st.font.name = "Liberation Sans"
st.font.size = Pt(9)
doc.add_heading("TRIPOD+AI checklist - manuscript v5", level=0)
t = doc.add_table(rows=1, cols=len(header))
t.style = "Table Grid"
for j, h in enumerate(header):
    c = t.cell(0, j)
    c.text = h
    for p in c.paragraphs:
        for run in p.runs:
            run.font.bold = True
            run.font.size = Pt(8)
for r in data:
    cells = t.add_row().cells
    for j, val in enumerate(r):
        cells[j].text = val
        for p in cells[j].paragraphs:
            for run in p.runs:
                run.font.size = Pt(7.5)
doc.save(DOCX)
print("saved", DOCX)
