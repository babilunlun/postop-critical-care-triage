#!/usr/bin/env python3
"""Build manuscript_v4.docx from manuscript_v4.md + v4 table CSVs (python-docx).
v4: ICU-course-verified triage repositioning; embeds Tables 1-4."""
import csv
import re
from docx import Document
from docx.shared import Pt

MD = "/mnt/results/04_manuscript/manuscript_v4.md"
T1 = "/mnt/results/04_manuscript/table1_combined_v4.csv"
T2 = "/mnt/results/04_manuscript/table2_performance_v4.csv"
T3 = "/mnt/results/06_icu_course/table3_outcome_decomposition.csv"
T4 = "/mnt/results/06_icu_course/table4_triage_pathway.csv"
OUT = "/workspace/manuscript_v4.docx"

doc = Document()

style = doc.styles["Normal"]
style.font.name = "Liberation Sans"
style.font.size = Pt(11)
for lvl in ("Heading 1", "Heading 2"):
    st = doc.styles[lvl]
    st.font.name = "Liberation Sans"

text = open(MD, encoding="utf-8").read()
lines = text.split("\n")


def add_table_from_csv(path, doc, font_size=8):
    with open(path, encoding="utf-8") as f:
        rows = list(csv.reader(f))
    t = doc.add_table(rows=len(rows), cols=len(rows[0]))
    t.style = "Table Grid"
    t.autofit = True
    for i, r in enumerate(rows):
        for j, val in enumerate(r):
            cell = t.cell(i, j)
            cell.text = val.strip()
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(font_size)
                    run.font.name = "Liberation Sans"
                    if i == 0:
                        run.font.bold = True
    return t


def clean_caption(ln):
    c = re.sub(r"\s*\[Full table:[^\]]*\]", "", ln)
    return c.replace("**", "").replace("*", "").replace("`", "").strip()


i = 0
in_refs = False
in_tables = False
while i < len(lines):
    ln = lines[i].rstrip()
    i += 1
    if not ln.strip() or ln.strip() == "---":
        continue
    if ln.startswith("# "):
        doc.add_heading(ln[2:].strip(), level=0)
        continue
    if ln.startswith("## "):
        h = ln[3:].strip()
        doc.add_heading(h, level=1)
        in_refs = (h == "References")
        in_tables = (h == "Tables")
        continue
    if ln.startswith("### "):
        doc.add_heading(ln[4:].strip(), level=2)
        continue
    if in_tables and ln.startswith("**Table 1"):
        doc.add_paragraph(clean_caption(ln))
        add_table_from_csv(T1, doc, font_size=6)
        doc.add_paragraph("")
        continue
    if in_tables and ln.startswith("**Table 2"):
        doc.add_paragraph(clean_caption(ln))
        add_table_from_csv(T2, doc, font_size=8)
        doc.add_paragraph("")
        continue
    if in_tables and ln.startswith("**Table 3"):
        doc.add_paragraph(clean_caption(ln))
        add_table_from_csv(T3, doc, font_size=7)
        doc.add_paragraph("")
        continue
    if in_tables and ln.startswith("**Table 4"):
        doc.add_paragraph(clean_caption(ln))
        add_table_from_csv(T4, doc, font_size=7)
        doc.add_paragraph("")
        continue
    if ln.startswith("- "):
        doc.add_paragraph(ln[2:].strip(), style="List Bullet")
        continue
    m = re.match(r"^(\d+)\.\s+(.*)$", ln)
    if in_refs and m:
        clean = m.group(2).replace("**", "").replace("*", "").replace("`", "")
        doc.add_paragraph(f"{m.group(1)}. {clean}")
        continue
    clean = ln.replace("**", "").replace("*", "").replace("`", "").replace("^", "")
    doc.add_paragraph(clean)

doc.save(OUT)
print("saved", OUT)

from docx import Document as _D
d2 = _D(OUT)
print("paragraphs:", len(d2.paragraphs), "tables:", len(d2.tables))
for k, tb in enumerate(d2.tables):
    print(f"  table {k+1}: {len(tb.rows)} rows x {len(tb.columns)} cols")
