#!/usr/bin/env python3
"""Build tripod_ai_checklist_v4.csv/.docx from v3 with ICU-course/triage repositioning updates."""
import csv
from docx import Document
from docx.shared import Pt

SRC = "/mnt/results/04_manuscript/tripod_ai_checklist_v3.csv"
CSV = "/mnt/results/04_manuscript/tripod_ai_checklist_v4.csv"

rows = list(csv.reader(open(SRC, encoding="utf-8")))
header, data = rows[0], rows[1:]

# item -> (Reported-in-manuscript location, Evidence/notes)
updates = {
    "1": ("Title",
          "Title identifies development AND external validation of a multivariable prediction model, "
          "the target population (non-cardiac surgery), the outcome (postoperative critical care), and "
          "the ICU-course-verified triage evaluation."),
    "2": ("Abstract",
          "139-word abstract covers setting, sample, model, outcome, internal/temporal/international "
          "performance, the ICU-course decomposition (57.5%/72.8% observational admissions; "
          "admitted-patient discrimination AUROC 0.50, with the three-variable SASA significantly "
          "better on that contrast), and the need-recalibrated three-tier triage pathway."),
    "3b": ("Introduction, para 1 and 3; Discussion - Clinical positioning",
           "Intended use restated as an end-of-surgery triage aid (not an admission oracle): low tier "
           "supports safe de-escalation, high tier supports proactive ICU planning, intermediate tier "
           "flags patients for heightened human assessment."),
    "4": ("Introduction - Objectives",
          "Objectives: (i) develop the model; (ii) validate the frozen model in two independent cohorts "
          "(INSPIRE temporal; MOVER international); (iii) head-to-head GRU vs XGBoost; (iv) using "
          "ICU-course data, separate true critical care need from admission behavior and derive a "
          "risk-stratified triage pathway."),
    "6b": ("Methods - Participants",
           "Identical clinical gates across cohorts; MOVER de-duplicated by encounter identifier "
           "(49,394 rows -> 48,370 operations in 34,143 patients; 1,024 duplicate rows, outcomes 100% "
           "concordant)."),
    "7": ("Methods - Predictors; ICU-course verification",
          "Predictor harmonization as before; NEW: ICU-course extraction described (INSPIRE bedside "
          "flowsheet + medication + procedure tables, 99.1% coverage; MOVER medication administration "
          "records + airway documentation, quality gate 93.6%; ICD-10-PCS for sensitivity only)."),
    "8a": ("Methods - Outcome; ICU-course-verified outcome categories",
           "Composite of early postoperative ICU admission or in-hospital death, defined per cohort "
           "(INSPIRE: ICU admission <=24 h after anesthesia end and before discharge, deaths counted "
           "to 24 h after discharge; MOVER: encounter-level ICU flag). NEW: four mutually exclusive "
           "ICU-course-verified categories (ICU-dependent, observational, missed escalation, "
           "uncomplicated ward) and the any-true-need union, with cohort-specific definitions "
           "disclosed."),
    "12d": ("Methods - External validation; Triage pathway evaluation; Results",
            "Three-cohort design; performance heterogeneity quantified across internal, temporal, and "
            "international validation; NEW: refined-outcome AUROCs (composite, ICU-dependent, "
            "any-true-need, observational-vs-ward, ICU-dependent-vs-observational) and within-band "
            "discrimination in the intermediate tier; all SASA-model comparisons are paired "
            "complete-case analyses on the SASA-evaluable subset with DeLong tests for correlated "
            "ROC curves."),
    "12f": ("Methods - External validation and recalibration; Triage pathway evaluation",
            "Intercept-only and Platt recalibration with 500x 50/50 split-sample validation in both "
            "cohorts; NEW: 10-fold cross-validated Platt recalibration of frozen predictions against "
            "the any-true-need outcome (need-recalibrated risk), with the SASA recalibrated to the "
            "same target as comparator."),
    "15": ("Methods - Triage pathway evaluation; Results - A risk-stratified triage pathway",
           "Model output is a probability; for triage deployment, pre-specified need-recalibrated risk "
           "bands (low <5%, intermediate 5-30%, high >30%) define the three-tier pathway, with "
           "percentile-band robustness checks."),
    "16": ("Methods - Data sources, Predictors, Outcome, Cohort overlap; Discussion - Limitations",
           "Development-vs-evaluation differences documented, including outcome windows (VitalDB ICU "
           "stay >=1 day; INSPIRE 24-h window; MOVER encounter-level flag) and asymmetric ICU-course "
           "capture (ventilation-dominant INSPIRE flowsheet vs vasopressor-dominant MOVER MAR)."),
    "20a": ("Results - Cohort characteristics; Fig. 1; Methods - Participants, Cohort overlap",
            "Updated flow: INSPIRE 98,633 -> 2,437 development-linked exclusions -> 96,196 operations "
            "(78,305 patients; 10,370 events, 10.8%); MOVER 49,394 rows -> 1,024 duplicates removed -> "
            "48,370 operations (34,143 patients; 21,740 events, 44.9%)."),
    "23a": ("Tables 2-4; Figures 2, 4, 5; Supplementary Tables S1-S3, S5, S7-S11",
            "AUROC/AUPRC/Brier/calibration with DeLong CIs for internal, INSPIRE (XGBoost AUROC 0.907, "
            "0.905-0.910), and de-duplicated MOVER (0.794, 0.790-0.798); refined-outcome AUROCs "
            "(ICU-dependent 0.877/0.783; admitted-patient contrast 0.502/0.640, with paired "
            "complete-case SASA comparisons on the SASA-evaluable subset); triage-tier event "
            "rates in Table 4 (SASA rows on the SASA-evaluable denominator)."),
    "23b": ("Results - Temporal/External validation, ICU-course-verified refinement; Discussion",
            "Discrimination gradient: internal 0.932 -> temporal (INSPIRE) 0.907 -> international "
            "(MOVER) 0.794; calibration drift moderate temporally (intercept -0.81, slope 0.88), large "
            "internationally (intercept 2.35)."),
    "24": ("Results - Temporal/External validation, Triage pathway; Tables 2 and 4",
           "Platt recalibration restored calibration in INSPIRE (split-sample Brier 0.066, slope "
           "0.999) and MOVER (Brier 0.183, slope 1.001); need-recalibrated tiers showed consistent "
           "safety across cohorts (low-tier mortality <0.1% and NPV ~98% in both); band event rates "
           "match band definitions by construction under within-cohort recalibration, while tier "
           "sizes differed with case mix (low tier 74.5% vs 32.0%)."),
    "25": ("Discussion, para 1-7",
           "Interpretation updated: discrimination transfers while absolute risk requires local "
           "recalibration; ICU-course verification shows the model predicts admission behavior as "
           "much as need (admission is a clinician decision, refs 23-25), motivating triage-aid "
           "deployment with the uncertain middle flagged for human assessment; model-choice "
           "reversal disclosed (LR full significantly better temporally, XGBoost better internally "
           "and internationally; differences <=0.04 AUROC)."),
    "26": ("Discussion - Limitations",
           "Updated limitations: behavioral component of the composite outcome; observational "
           "admission not equal to unnecessary admission; asymmetric intervention capture across "
           "cohorts; MOVER lacks ICU timestamps (missed escalations death-only) and PCS codes are "
           "patient-level; post-hoc outcome refinement with frozen models; residual overlap bound "
           "+0.0006 AUROC; pre-deduplication supplementary analyses (impact 0.0002 AUROC); SASA "
           "comparisons are complete-case (estimated-blood-loss missingness 40.3%/38.8%, "
           "informative), mitigated by paired SASA-evaluable analyses; triage band event rates "
           "approximate band definitions by construction under within-cohort recalibration."),
    "27b": ("Discussion - Clinical positioning",
            "User interaction defined per tier: low tier supports de-escalation, high tier proactive "
            "ICU planning, intermediate tier triggers heightened human assessment (senior review, "
            "extended recovery, monitored ward); the model must not be used to deny ICU admission."),
    "27c": ("Discussion - Clinical positioning; Conclusion",
            "Next steps: silent-mode prospective evaluation and workflow-integration studies; the "
            "deployable unit is a locally recalibrated risk of true need, not the raw score."),
}

n_applied = 0
for r in data:
    item = r[1]
    if item in updates:
        r[3], r[4] = updates[item]
        n_applied += 1

assert n_applied == len(updates), f"applied {n_applied} of {len(updates)} updates (item id mismatch?)"

with open(CSV, "w", encoding="utf-8", newline="") as f:
    w = csv.writer(f)
    w.writerow(header)
    w.writerows(data)
print("CSV saved:", CSV, "rows:", len(data), "updates applied:", n_applied)

doc = Document()
st = doc.styles["Normal"]
st.font.name = "Liberation Sans"
st.font.size = Pt(9)
doc.add_heading("TRIPOD+AI 2024 checklist - manuscript v4 (ICU-course-verified triage evaluation)", level=1)
t = doc.add_table(rows=len(data) + 1, cols=5)
t.style = "Table Grid"
for j, h in enumerate(header):
    t.cell(0, j).text = h
for i, r in enumerate(data, start=1):
    for j, val in enumerate(r):
        t.cell(i, j).text = val
for ri, row in enumerate(t.rows):
    for cell in row.cells:
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.size = Pt(8)
                run.font.name = "Liberation Sans"
                if ri == 0:
                    run.font.bold = True
doc.save("/workspace/tripod_ai_checklist_v4.docx")
print("DOCX saved")
