#!/usr/bin/env python3
"""Build tripod_ai_checklist_v2.csv/.docx from v1 with sensitivity-analysis updates."""
import csv
from docx import Document
from docx.shared import Pt

SRC = "/mnt/results/04_manuscript/tripod_ai_checklist_v1.csv"
CSV = "/mnt/results/04_manuscript/tripod_ai_checklist_v2.csv"

rows = list(csv.reader(open(SRC, encoding="utf-8")))
header, data = rows[0], rows[1:]

updates = {
    "2": ("Abstract",
          "149-word unstructured abstract covers setting, sample, model, outcome, internal/external performance, recalibration, DL comparator, outcome-definition sensitivity, conclusion."),
    "12d": ("Methods - External validation; Results - Subgroup and Sensitivity analyses",
            "Two-centre design; performance heterogeneity quantified by external validation, subgroup analyses (department, year, patient class), and outcome-definition sensitivity subsets (routine-ICU exclusion; ambulatory-only)."),
    "12e": ("Methods - Statistical analysis",
            "AUROC (DeLong CI), AUPRC, Brier, calibration intercept/slope, decision curves (dcurves); DeLong test for model comparison."),
    "12f": ("Methods - External validation and recalibration; Methods - Sensitivity analyses",
            "Intercept-only and Platt recalibration; 500x 50/50 split-sample validation; recalibration refit within each sensitivity subset."),
    "23a": ("Tables 2; Figures 2-4; Supplementary Figures S2-S3; Supplementary Tables S1-S3, S5",
            "AUROC/AUPRC/Brier/calibration with DeLong CIs; subgroup AUROCs; sensitivity-subset metrics; ROC/calibration/decision-curve plots."),
    "24": ("Results - External validation and Sensitivity analyses; Table 2; Figure 4; Supplementary Table S5",
           "Intercept-only and Platt recalibration with split-sample-validated Brier and calibration slope, overall and within sensitivity subsets."),
    "25": ("Discussion, para 1-5",
           "Findings interpreted against prior single-center and DL studies; discrimination-vs-calibration message; outcome-definition sensitivity and arterial-line ablation interpreted."),
    "27c": ("Discussion, para 6-7",
            "Silent-mode prospective evaluation; local recalibration set; further validation cohorts."),
}

for r in data:
    item = r[1]
    if item in updates:
        r[3], r[4] = updates[item]

with open(CSV, "w", encoding="utf-8", newline="") as f:
    w = csv.writer(f)
    w.writerow(header)
    w.writerows(data)
print("CSV saved:", CSV, "rows:", len(data))

doc = Document()
st = doc.styles["Normal"]
st.font.name = "Liberation Sans"
st.font.size = Pt(9)
doc.add_heading("TRIPOD+AI 2024 checklist - manuscript v2", level=1)
t = doc.add_table(rows=len(data) + 1, cols=5)
t.style = "Table Grid"
for j, h in enumerate(header):
    t.cell(0, j).text = h
for i, r in enumerate(data, start=1):
    for j, val in enumerate(r):
        t.cell(i, j).text = val
for row in t.rows:
    for cell in row.cells:
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.size = Pt(8)
                run.font.name = "Liberation Sans"
doc.save("/workspace/tripod_ai_checklist_v2.docx")
print("DOCX saved")
